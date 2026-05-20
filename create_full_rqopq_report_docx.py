#!/usr/bin/env python3
"""Create a DOCX report for the full multimodal RQ-OPQ semantic SID run."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(".")
SID_DIR = ROOT / "multimodal_rqopq_full_semantic"
OUT = ROOT / "full_rqopq_semantic_report.docx"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def pct(x: float, digits: int = 2) -> str:
    return f"{x * 100:.{digits}f}%"


def fmt_int(x: int | float) -> str:
    return f"{int(x):,}"


def fmt_float(x: float, digits: int = 4) -> str:
    return f"{x:.{digits}f}"


def set_run_font(run, size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, size=11, color="222222", after=6, line_spacing=1.10) -> None:
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B7C3D1") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        existing = borders.find(qn(f"w:{edge}"))
        if existing is not None:
            borders.remove(existing)
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)


def set_table_widths(table, widths_in: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in, strict=True):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_table(table, widths_in: list[float], header_fill="F2F4F7") -> None:
    set_table_widths(table, widths_in)
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=9,
                        bold=(row_idx == 0),
                        color="1F3349" if row_idx == 0 else "222222",
                    )


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, color="2E74B5" if level <= 2 else "1F4D78")


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    style_paragraph(paragraph)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text, style="List Bullet")
    style_paragraph(paragraph, after=4, line_spacing=1.167)


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths=(2.35, 4.15)) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "值"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    style_table(table, list(widths))
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            cell.text = value
    style_table(table, widths)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    set_run_font(run, size=11, bold=True, color="1F3A5F")
    paragraph.add_run("  ")
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color="222222")
    style_paragraph(paragraph, after=0, line_spacing=1.15)
    set_table_borders(table, color="D8DEE8")
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")

    for name, size, color in [
        ("Title", 22, "0B2545"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Full RQ-OPQ Semantic SID Report")
    set_run_font(run, size=9, color="536173")


def main() -> None:
    metrics = load_json(SID_DIR / "metrics.json")
    analysis = load_json(SID_DIR / "analysis_summary.json")

    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("全量商品 RQ-OPQ Semantic SID 实验报告")
    style_paragraph(title, size=22, color="0B2545", after=4, line_spacing=1.05)
    subtitle = doc.add_paragraph(
        f"实验目录：{SID_DIR.as_posix()} | 生成日期：{date.today().isoformat()}"
    )
    style_paragraph(subtitle, size=10, color="536173", after=14)

    add_callout(
        doc,
        "结论摘要",
        "全量 semantic_mean 结果整体可用：量化质量稳定、codebook 使用健康；主要风险是 collision 从 100w 样本的 0.84% 上升到全量 2.76%，且存在少数超大 collision group，需要结合内部商品页抽样确认是否为同款/模板化商品。",
    )

    cfg = metrics["config"]
    add_heading(doc, "1. 实验配置", 1)
    add_kv_table(
        doc,
        [
            ("输入数据", "title_emb.npy, image_emb.npy, item_feat.npz, itemid.npy"),
            ("商品数", fmt_int(metrics["data"]["num_items"])),
            ("融合特征维度", str(metrics["fusion"]["fused_dim"])),
            ("离散特征映射", metrics["fusion"]["discrete_embedding_mode"]),
            ("权重", "title=1.0, image=1.0, cat=0.5, label=0.5, flag=0.2"),
            ("RQ 参数", f"K={cfg['rq_clusters']}, L={cfg['rq_levels']}, code_dim={cfg['code_dim']}"),
            ("OPQ 参数", f"M={cfg['opq_subspaces']}, K={cfg['opq_clusters']}, subdim={cfg['code_dim'] // cfg['opq_subspaces']}"),
            ("SID token 数", str(metrics["sid"]["categorical_tokens_per_item"])),
            ("如保存 dense concat 维度", str(metrics["sid"]["dense_concat_dim_if_saved"])),
            ("本次 dense 文件", "未保存 sid_codeword_concat / sid_reconstruction"),
        ],
    )

    add_heading(doc, "2. 核心结果", 1)
    col = metrics["collisions"]
    rec = analysis["sampled_reconstruction"]
    add_table(
        doc,
        ["指标", "结果", "解释"],
        [
            ["unique SID", fmt_int(col["unique_sid"]), "最终 5-token SID 的唯一组合数"],
            ["unique SID ratio", pct(col["unique_sid_ratio"]), "unique SID / item 数"],
            ["extra collision rate", pct(col["collision_rate_extra"], 3), "(item 数 - unique SID) / item 数"],
            ["collision item 覆盖率", pct(analysis["collision_distribution"]["collision_items"] / analysis["num_items"], 3), "落在非 singleton group 中的 item 占比"],
            ["max collision group", fmt_int(col["max_collision_group_size"]), "最大 SID group 中 item 数"],
            ["PCA explained variance", pct(metrics["pca"]["explained_variance_ratio_sum"], 2), "32 维 PCA 保留的方差占比"],
            ["OPQ final SSE", fmt_float(metrics["opq"]["final_original_space_mse"], 5), "每个 item 的 32 维 residual 平方误差总和"],
            ["sampled cosine mean", fmt_float(rec["cosine_mean"], 4), "10w 抽样中 reconstructed vector 与 projected vector 的平均 cosine"],
        ],
        [2.0, 1.45, 3.05],
    )

    add_heading(doc, "3. 量化质量", 1)
    rq_rows = []
    for i, level in enumerate(metrics["rq_kmeans"]["per_level"], start=1):
        rq_rows.append(
            [
                f"RQ L{i}",
                fmt_float(level["residual_mse_after_level"], 5),
                f"{level['used_codes']}/512",
                fmt_float(level["entropy"], 4),
                f"{fmt_int(level['min_cluster_size'])} / {fmt_int(level['max_cluster_size'])}",
            ]
        )
    rq_rows.append(
        [
            "OPQ final",
            fmt_float(metrics["opq"]["final_original_space_mse"], 5),
            "2 x 256",
            "约 7.997",
            "见 OPQ 子空间统计",
        ]
    )
    add_table(
        doc,
        ["阶段", "Residual SSE", "码字使用", "Entropy", "Min / Max group"],
        rq_rows,
        [1.1, 1.35, 1.2, 1.15, 1.7],
    )
    add_body(
        doc,
        "RQ 残差从 L1 的 0.14116 降到 L3 的 0.06649；OPQ 进一步降到 0.03892，说明 OPQ 对最后 residual 的子空间量化确实带来了额外压缩收益。抽样重建 cosine mean 为 0.9492，p50 为 0.9550，p05 仍约 0.9000，整体保真度较好。",
    )

    add_heading(doc, "4. Collision 与分层 CUR", 1)
    prefix_rows = []
    capacities = [512, 512, 512, 256, 256]
    cap = 1
    for stat, c in zip(analysis["prefix_stats"], capacities, strict=True):
        cap *= c
        prefix_rows.append(
            [
                str(stat["depth"]),
                fmt_int(stat["unique"]),
                fmt_int(cap),
                pct(stat["unique"] / cap, 4),
                pct(stat["unique_ratio"], 2),
                fmt_int(stat["max_group_size"]),
            ]
        )
    add_table(
        doc,
        ["深度", "Used prefix", "Total prefix space", "分层 CUR", "Unique/item", "Max group"],
        prefix_rows,
        [0.55, 1.15, 1.55, 1.05, 1.05, 1.15],
    )
    add_body(
        doc,
        "分层 CUR 的定义是 used SID prefix / 理论 prefix 空间。前两层组合空间利用率很高，depth=2 达到 99.11%；从 depth=3 开始理论空间远大于商品数，因此 CUR 下降是正常现象。更有业务解释力的是 unique/item：depth=3 已经区分 72.68% item，加入两个 OPQ token 后提升到 97.24%。",
    )

    dist = analysis["collision_distribution"]
    add_table(
        doc,
        ["Collision 分布", "值"],
        [
            ["collision groups", fmt_int(dist["collision_groups"])],
            ["collision items", fmt_int(dist["collision_items"])],
            ["mean group size", fmt_float(dist["mean_collision_group_size"], 2)],
            ["p50 group size", fmt_float(dist["quantiles"]["0.5"], 0)],
            ["p90 group size", fmt_float(dist["quantiles"]["0.9"], 0)],
            ["p99 group size", fmt_float(dist["quantiles"]["0.99"], 0)],
            ["p99.9 group size", fmt_float(dist["quantiles"]["0.999"], 1)],
            ["max group size", fmt_int(dist["max_group_size"])],
        ],
        [2.6, 3.9],
    )
    add_body(
        doc,
        "大多数 collision group 很小：p50 为 2，p90 为 3，p99 为 9。风险主要集中在 top collision group，例如最大 SID 235-296-296-212-140 包含 2,354 个 item。建议优先抽样检查 top 20 group 是否为高度相似、模板化或近重复商品。",
    )

    add_heading(doc, "5. 同款不同颜色 Case Study", 1)
    add_body(
        doc,
        "抽查 item 103189、103192、103191、103196、103194、103187、103181、103183、103184、103186 后发现：这些 item 的 cat_ids 完全一致，flags 完全一致，labels 也高度相似，但 SID prefix 可能差异较大。",
    )
    add_table(
        doc,
        ["item_id", "SID", "与 103189 的 title cos", "image cos", "projected cos"],
        [
            ["103189", "457-112-397-195-67", "1.000", "1.000", "1.000"],
            ["103192", "457-112-315-122-222", "0.926", "0.800", "0.842"],
            ["103191", "457-424-163-4-131", "0.951", "0.804", "0.819"],
            ["103196", "457-112-146-62-161", "0.959", "0.889", "0.940"],
            ["103194", "241-412-381-4-45", "0.966", "0.865", "0.960"],
            ["103187", "501-112-446-54-3", "0.887", "0.689", "0.683"],
            ["103181", "203-98-450-4-248", "0.909", "0.461", "0.658"],
            ["103183", "457-112-86-213-194", "0.965", "0.809", "0.886"],
            ["103184", "104-38-381-187-81", "0.925", "0.820", "0.748"],
            ["103186", "170-424-390-204-165", "0.886", "0.597", "0.629"],
        ],
        [0.85, 1.65, 1.2, 1.0, 1.1],
    )
    add_body(
        doc,
        "这个现象不是映射错误，更像是当前无监督 RQ-OPQ 目标与业务期望不完全一致：当前 SID 压缩的是融合 embedding，而不是显式建模 SPU/同款关系。染发剂图片中颜色是强视觉信号，标题中也包含颜色词，因此同款不同颜色会被 image/title embedding 拉开。KMeans 又是硬分配，即使两个 item 在 projected 空间很近，也可能落到不同的第一层 cell。",
    )
    add_callout(
        doc,
        "解释重点",
        "当前结果适合作为 item-level 语义压缩特征；如果业务要求 prefix 表达同款/SPU 层级，需要引入 parent_item_id/SPU/variation group、去颜色 title embedding，或用 group mean embedding 训练前几层 SID。",
    )

    add_heading(doc, "6. 使用建议与后续实验", 1)
    add_bullet(doc, "当前全量 semantic_mean 结果可以先用于下游特征实验，尤其适合作为 5 个 categorical token 输入模型。")
    add_bullet(doc, "不要仅用 prefix 是否相同判断商品是否相似；对当前方案，应同时参考 projected cosine、reconstructed cosine 或 codeword concat cosine。")
    add_bullet(doc, "如果 top collision group 在内部查询后确认为相似/近重复商品，则 2.76% extra collision rate 可以接受。")
    add_bullet(doc, "如果希望同款不同颜色共享更稳定的 prefix，优先尝试 RQ L=4 或降低 image 权重；更理想的是引入 SPU/variation group 做 family-aware SID。")
    add_bullet(doc, "全量 collision 查看表已生成在 multimodal_rqopq_full_semantic/sid_item_index/，其中 sid_groups_summary.csv 适合按 group size 抽查。")

    add_heading(doc, "7. 关键产物", 1)
    add_kv_table(
        doc,
        [
            ("metrics.json", "训练配置、PCA、RQ/OPQ、collision 核心指标"),
            ("analysis_summary.json", "补算的 prefix CUR、collision 分布、抽样重建质量、近邻 prefix agreement"),
            ("sid_codes.npy", "未 offset 的 5-token SID"),
            ("sid_codes_offset.npy", "适合单 embedding table 使用的 offset token"),
            ("sid_item_index/sid_groups_summary.csv", "按 SID group 聚合的 item_id 索引"),
            ("sid_item_index/collision_items.csv", "collision group 中每个 item 的明细"),
        ],
    )

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
