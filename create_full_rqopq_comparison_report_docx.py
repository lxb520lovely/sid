#!/usr/bin/env python3
"""Create a DOCX comparison report for full random vs semantic_mean RQ-OPQ runs."""

from __future__ import annotations

import json
import math
from datetime import date
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(".")
RUNS = {
    "random": ROOT / "multimodal_rqopq_full_random",
    "semantic_mean": ROOT / "multimodal_rqopq_full_semantic",
}
FIG_DIR = ROOT / "full_rqopq_comparison_figures"
OUT = ROOT / "full_rqopq_random_vs_semantic_report.docx"


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def pct(x: float, digits: int = 2) -> str:
    return f"{x * 100:.{digits}f}%"


def fmt_int(x: int | float) -> str:
    return f"{int(x):,}"


def fmt_float(x: float, digits: int = 4) -> str:
    return f"{x:.{digits}f}"


def rec_value(rec: dict[str, Any], key: str) -> float:
    aliases = {
        "per_dim_mse_mean": ["per_dim_mse_mean", "mse_mean"],
        "per_dim_mse_p50": ["per_dim_mse_p50", "mse_p50"],
        "per_dim_mse_p95": ["per_dim_mse_p95", "mse_p95"],
        "per_dim_mse_p99": ["per_dim_mse_p99", "mse_p99"],
    }
    for candidate in aliases.get(key, [key]):
        if candidate in rec:
            return float(rec[candidate])
    raise KeyError(key)


def set_run_font(run, size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, size=11, color="222222", after=6, line_spacing=1.10) -> None:
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B7C3D1") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        existing = borders.find(qn(f"w:{edge}"))
        if existing is not None:
            borders.remove(existing)
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)


def style_table(table, widths_in: list[float], header_fill="F2F4F7") -> None:
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell, width in zip(row.cells, widths_in, strict=True):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=8.7,
                        bold=(row_idx == 0),
                        color="1F3349" if row_idx == 0 else "222222",
                    )
    set_table_borders(table)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    for name, size, color in [
        ("Title", 21, "0B2545"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Full RQ-OPQ Random vs Semantic Report")
    set_run_font(run, size=9, color="536173")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, color="2E74B5" if level <= 2 else "1F4D78")


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    style_paragraph(paragraph)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text, style="List Bullet")
    style_paragraph(paragraph, after=4, line_spacing=1.167)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            cell.text = value
    style_table(table, widths)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    set_run_font(run, size=11, bold=True, color="1F3A5F")
    paragraph.add_run("  ")
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color="222222")
    style_paragraph(paragraph, after=0, line_spacing=1.15)
    set_table_borders(table, color="D8DEE8")
    doc.add_paragraph()


def add_figure(doc: Document, image: Path, caption: str, width=6.2) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_run_font(run, size=9, italic=True, color="536173")
    cap.paragraph_format.space_after = Pt(8)


def load_runs() -> dict[str, dict[str, Any]]:
    runs = {}
    for name, path in RUNS.items():
        runs[name] = {
            "dir": path,
            "metrics": load_json(path / "metrics.json"),
            "analysis": load_json(path / "analysis_summary.json"),
        }
    return runs


def make_figures(runs: dict[str, dict[str, Any]]) -> dict[str, Path]:
    FIG_DIR.mkdir(exist_ok=True)
    colors = {"random": "#4C78A8", "semantic_mean": "#F58518"}
    labels = {"random": "random", "semantic_mean": "semantic_mean"}
    figures: dict[str, Path] = {}

    plt.style.use("seaborn-v0_8-whitegrid")

    fig, ax = plt.subplots(figsize=(7.2, 4.2), dpi=180)
    stages = ["RQ1", "RQ2", "RQ3", "OPQ"]
    for name, run in runs.items():
        m = run["metrics"]
        y = [x["residual_mse_after_level"] for x in m["rq_kmeans"]["per_level"]]
        y.append(m["opq"]["final_original_space_mse"])
        ax.plot(stages, y, marker="o", linewidth=2.2, label=labels[name], color=colors[name])
    ax.set_title("Residual SSE by Quantization Stage")
    ax.set_ylabel("SSE per item")
    ax.legend()
    figures["residual"] = FIG_DIR / "residual_sse.png"
    fig.tight_layout()
    fig.savefig(figures["residual"])
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(7.2, 4.2), dpi=180)
    metrics = ["extra collision", "collision item", "unique SID"]
    x = np.arange(len(metrics))
    width = 0.34
    for offset, name in [(-width / 2, "random"), (width / 2, "semantic_mean")]:
        a = runs[name]["analysis"]
        m = runs[name]["metrics"]
        vals = [
            m["collisions"]["collision_rate_extra"] * 100,
            a["collision_distribution"]["collision_items"] / a["num_items"] * 100,
            m["collisions"]["unique_sid_ratio"] * 100,
        ]
        ax.bar(x + offset, vals, width=width, label=labels[name], color=colors[name])
    ax.set_xticks(x)
    ax.set_xticklabels(metrics)
    ax.set_ylabel("percent")
    ax.set_title("Uniqueness and Collision")
    ax.legend()
    figures["collision"] = FIG_DIR / "collision.png"
    fig.tight_layout()
    fig.savefig(figures["collision"])
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(7.2, 4.2), dpi=180)
    depths = [1, 2, 3, 4, 5]
    for name, run in runs.items():
        y = [x["unique_ratio"] * 100 for x in run["analysis"]["prefix_stats"]]
        ax.plot(depths, y, marker="o", linewidth=2.2, label=labels[name], color=colors[name])
    ax.set_xticks(depths)
    ax.set_xlabel("SID prefix depth")
    ax.set_ylabel("unique prefix / item (%)")
    ax.set_title("Prefix Uniqueness Curve")
    ax.legend()
    figures["prefix"] = FIG_DIR / "prefix_unique.png"
    fig.tight_layout()
    fig.savefig(figures["prefix"])
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(7.2, 4.2), dpi=180)
    capacities = [512, 512, 512, 256, 256]
    for name, run in runs.items():
        cap = 1
        y = []
        for stat, c in zip(run["analysis"]["prefix_stats"], capacities, strict=True):
            cap *= c
            y.append(stat["unique"] / cap * 100)
        ax.plot(depths, y, marker="o", linewidth=2.2, label=labels[name], color=colors[name])
    ax.set_yscale("log")
    ax.set_xticks(depths)
    ax.set_xlabel("SID prefix depth")
    ax.set_ylabel("hierarchical CUR (%) log scale")
    ax.set_title("Hierarchical Code Usage Rate")
    ax.legend()
    figures["cur"] = FIG_DIR / "hierarchical_cur.png"
    fig.tight_layout()
    fig.savefig(figures["cur"])
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(7.2, 4.2), dpi=180)
    components = ["RQ1", "RQ2", "RQ3", "OPQ0", "OPQ1"]
    x = np.arange(len(components))
    for offset, name in [(-width / 2, "random"), (width / 2, "semantic_mean")]:
        m = runs[name]["metrics"]
        ratios = [level["entropy"] / math.log2(512) for level in m["rq_kmeans"]["per_level"]]
        ratios += [sub["entropy"] / math.log2(256) for sub in m["opq"]["history"][-1]["per_subspace"]]
        ax.bar(x + offset, [v * 100 for v in ratios], width=width, label=labels[name], color=colors[name])
    ax.set_ylim(94, 100.5)
    ax.set_xticks(x)
    ax.set_xticklabels(components)
    ax.set_ylabel("entropy / max entropy (%)")
    ax.set_title("Codebook Usage Balance")
    ax.legend()
    figures["entropy"] = FIG_DIR / "entropy.png"
    fig.tight_layout()
    fig.savefig(figures["entropy"])
    plt.close(fig)

    return figures


def comparison_rows(runs: dict[str, dict[str, Any]]) -> list[list[str]]:
    r_m = runs["random"]["metrics"]
    r_a = runs["random"]["analysis"]
    s_m = runs["semantic_mean"]["metrics"]
    s_a = runs["semantic_mean"]["analysis"]
    return [
        ["运行耗时", f"{r_m['elapsed_seconds']/60:.1f} min", f"{s_m['elapsed_seconds']/60:.1f} min", "random 更快；semantic_mean 多一次全量语义均值聚合"],
        ["PCA explained variance", pct(r_m["pca"]["explained_variance_ratio_sum"]), pct(s_m["pca"]["explained_variance_ratio_sum"]), "semantic_mean 略高"],
        ["RQ L3 residual SSE", fmt_float(r_m["rq_kmeans"]["per_level"][-1]["residual_mse_after_level"], 5), fmt_float(s_m["rq_kmeans"]["per_level"][-1]["residual_mse_after_level"], 5), "semantic_mean 低约 26.3%"],
        ["OPQ final SSE", fmt_float(r_m["opq"]["final_original_space_mse"], 5), fmt_float(s_m["opq"]["final_original_space_mse"], 5), "semantic_mean 低约 26.3%"],
        ["sampled cosine mean", fmt_float(r_a["sampled_reconstruction"]["cosine_mean"], 4), fmt_float(s_a["sampled_reconstruction"]["cosine_mean"], 4), "semantic_mean 略高"],
        ["unique SID ratio", pct(r_m["collisions"]["unique_sid_ratio"]), pct(s_m["collisions"]["unique_sid_ratio"]), "random 更高"],
        ["extra collision rate", pct(r_m["collisions"]["collision_rate_extra"], 3), pct(s_m["collisions"]["collision_rate_extra"], 3), "random 约为 semantic_mean 的一半"],
        ["collision item coverage", pct(r_a["collision_distribution"]["collision_items"] / r_a["num_items"], 3), pct(s_a["collision_distribution"]["collision_items"] / s_a["num_items"], 3), "random 更低"],
        ["max collision group", fmt_int(r_m["collisions"]["max_collision_group_size"]), fmt_int(s_m["collisions"]["max_collision_group_size"]), "semantic_mean top group 更大"],
    ]


def add_metric_definitions(doc: Document) -> None:
    add_heading(doc, "2. 指标口径", 1)
    add_table(
        doc,
        ["指标", "计算方式", "含义"],
        [
            ["Codebook CUR", "used code / K", "单层码本是否有塌缩；本实验各层均为 1.0"],
            ["Entropy ratio", "entropy(counts) / log2(K)", "码字使用是否均匀，越接近 1 越好"],
            ["Effective ratio", "2^entropy / K", "按熵折算的有效码字比例"],
            ["Hierarchical CUR", "unique prefix / theoretical prefix space", "分层组合空间使用率，例如 depth=2 为 used(RQ1,RQ2)/512^2"],
            ["Unique prefix ratio", "unique prefix / item count", "某个 prefix 深度能区分多少 item"],
            ["Extra collision rate", "(N - unique SID) / N", "额外冲突数量占比"],
            ["Collision item coverage", "items in non-singleton SID groups / N", "有多少 item 处在 collision group 内"],
            ["Residual SSE", "sum((x - x_hat)^2) per item", "脚本里 RQ/OPQ 的重建误差口径"],
            ["Reconstruction cosine", "cos(projected, reconstructed)", "SID codebook 重建向量与 PCA projected 向量的一致性"],
            ["Neighbor prefix agreement", "近邻 pair 中 prefix 完全相同的比例", "SID prefix 是否和连续空间近邻一致；本实验只做抽样估计"],
        ],
        [1.45, 2.35, 2.7],
    )


def main() -> None:
    runs = load_runs()
    figures = make_figures(runs)

    doc = Document()
    setup_styles(doc)
    title = doc.add_paragraph(style="Title")
    title.add_run("全量商品 RQ-OPQ SID：random vs semantic_mean 对比实验报告")
    style_paragraph(title, size=21, color="0B2545", after=4, line_spacing=1.05)
    subtitle = doc.add_paragraph(
        f"实验目录：multimodal_rqopq_full_random / multimodal_rqopq_full_semantic | 生成日期：{date.today().isoformat()}"
    )
    style_paragraph(subtitle, size=10, color="536173", after=14)

    add_callout(
        doc,
        "结论摘要",
        "semantic_mean 在量化重建质量上明显更好，OPQ final SSE 比 random 低约 26%；random 在 SID 唯一性上明显更好，extra collision rate 约为 1.39%，semantic_mean 为 2.76%。如果当前优先做下游特征增益，建议 random 与 semantic_mean 都进入离线/线上 A/B；如果必须先选一个保守版本，random 更稳；如果更重视语义聚类与重建质量，semantic_mean 更有价值。",
    )

    add_heading(doc, "1. 实验背景与配置", 1)
    m = runs["random"]["metrics"]
    add_table(
        doc,
        ["项目", "配置"],
        [
            ["输入", "title_emb.npy + image_emb.npy + item_feat.npz + itemid.npy"],
            ["商品数", fmt_int(m["data"]["num_items"])],
            ["融合维度", str(m["fusion"]["fused_dim"])],
            ["共同权重", "title=1.0, image=1.0, cat=0.5, label=0.5, flag=0.2"],
            ["RQ", "K=512, L=3, code_dim=32"],
            ["OPQ", "M=2, K=256, subdim=16"],
            ["SID 长度", "5 tokens = 3 RQ tokens + 2 OPQ tokens"],
            ["dense concat 维度", "128，如保存则为 3*32 + 2*16；本次未保存 dense 文件"],
            ["差异变量", "仅 discrete-embedding-mode 不同：random vs semantic_mean"],
        ],
        [2.0, 4.5],
    )
    add_body(
        doc,
        "random 模式使用确定性随机表把 cat_ids/labels 映射到 dense 向量；semantic_mean 模式先用 title/image 语义基向量统计每个 category/label 的均值，再作为离散特征 embedding。两者下游 RQ-OPQ 训练流程相同。",
    )

    add_metric_definitions(doc)

    add_heading(doc, "3. 总体对比", 1)
    add_table(
        doc,
        ["指标", "random", "semantic_mean", "解读"],
        comparison_rows(runs),
        [1.55, 1.2, 1.35, 2.4],
    )
    add_figure(doc, figures["residual"], "图 1：RQ/OPQ 各阶段 residual SSE。semantic_mean 在所有阶段误差更低。")
    add_figure(doc, figures["collision"], "图 2：唯一性与 collision 对比。random 的 SID 唯一性更强。")

    add_heading(doc, "4. 量化质量与重建质量", 1)
    rows = []
    for stage_idx in range(3):
        rows.append(
            [
                f"RQ L{stage_idx + 1}",
                fmt_float(runs["random"]["metrics"]["rq_kmeans"]["per_level"][stage_idx]["residual_mse_after_level"], 5),
                fmt_float(runs["semantic_mean"]["metrics"]["rq_kmeans"]["per_level"][stage_idx]["residual_mse_after_level"], 5),
                pct(1 - runs["semantic_mean"]["metrics"]["rq_kmeans"]["per_level"][stage_idx]["residual_mse_after_level"] / runs["random"]["metrics"]["rq_kmeans"]["per_level"][stage_idx]["residual_mse_after_level"], 1),
            ]
        )
    rows.append(
        [
            "OPQ final",
            fmt_float(runs["random"]["metrics"]["opq"]["final_original_space_mse"], 5),
            fmt_float(runs["semantic_mean"]["metrics"]["opq"]["final_original_space_mse"], 5),
            pct(1 - runs["semantic_mean"]["metrics"]["opq"]["final_original_space_mse"] / runs["random"]["metrics"]["opq"]["final_original_space_mse"], 1),
        ]
    )
    add_table(doc, ["阶段", "random SSE", "semantic_mean SSE", "semantic 优势"], rows, [1.25, 1.55, 1.7, 1.5])

    rec_rows = []
    for name in ["random", "semantic_mean"]:
        rec = runs[name]["analysis"]["sampled_reconstruction"]
        rec_rows.append(
            [
                name,
                fmt_float(rec_value(rec, "per_dim_mse_mean"), 6),
                fmt_float(rec.get("vector_sse_mean", rec_value(rec, "per_dim_mse_mean") * 32), 5),
                fmt_float(rec["cosine_mean"], 4),
                fmt_float(rec["cosine_p05"], 4),
                fmt_float(rec["cosine_p01"], 4),
            ]
        )
    add_table(
        doc,
        ["模式", "Per-dim MSE", "Vector SSE", "Cos mean", "Cos p05", "Cos p01"],
        rec_rows,
        [1.25, 1.2, 1.2, 1.05, 1.05, 1.05],
    )
    add_body(
        doc,
        "semantic_mean 的优势来自离散特征 embedding 本身具备语义均值结构，使得融合向量更容易被 RQ-OPQ 表达。random 的离散向量增加了区分度，但也引入更难压缩的随机方向，因此重建误差更高。",
    )

    add_heading(doc, "5. Codebook 使用情况", 1)
    usage_rows = []
    for name in ["random", "semantic_mean"]:
        m = runs[name]["metrics"]
        for idx, level in enumerate(m["rq_kmeans"]["per_level"], start=1):
            entropy_ratio = level["entropy"] / math.log2(512)
            effective_ratio = (2 ** level["entropy"]) / 512
            usage_rows.append(
                [
                    name,
                    f"RQ L{idx}",
                    f"{level['used_codes']}/512",
                    fmt_float(entropy_ratio, 4),
                    fmt_float(effective_ratio, 4),
                    f"{fmt_int(level['min_cluster_size'])} / {fmt_int(level['max_cluster_size'])}",
                ]
            )
        for sub in m["opq"]["history"][-1]["per_subspace"]:
            entropy_ratio = sub["entropy"] / math.log2(256)
            effective_ratio = (2 ** sub["entropy"]) / 256
            usage_rows.append(
                [
                    name,
                    f"OPQ {sub['subspace']}",
                    f"{sub['used_codes']}/256",
                    fmt_float(entropy_ratio, 4),
                    fmt_float(effective_ratio, 4),
                    f"{fmt_int(sub['min_cluster_size'])} / {fmt_int(sub['max_cluster_size'])}",
                ]
            )
    add_table(
        doc,
        ["模式", "组件", "CUR", "Entropy ratio", "Effective ratio", "Min / Max"],
        usage_rows,
        [1.1, 0.85, 0.95, 1.15, 1.15, 1.3],
    )
    add_figure(doc, figures["entropy"], "图 3：码字使用均匀性。两种模式都没有 codebook collapse。")

    add_heading(doc, "6. Collision、Prefix 与分层 CUR", 1)
    col_rows = []
    for name in ["random", "semantic_mean"]:
        m = runs[name]["metrics"]
        a = runs[name]["analysis"]
        dist = a["collision_distribution"]
        col_rows.append(
            [
                name,
                fmt_int(m["collisions"]["unique_sid"]),
                pct(m["collisions"]["unique_sid_ratio"]),
                pct(m["collisions"]["collision_rate_extra"], 3),
                pct(dist["collision_items"] / a["num_items"], 3),
                fmt_int(m["collisions"]["max_collision_group_size"]),
                fmt_float(dist["quantiles"]["0.99"], 0),
                fmt_float(dist["quantiles"]["0.999"], 1),
            ]
        )
    add_table(
        doc,
        ["模式", "Unique SID", "Unique ratio", "Extra collision", "Collision items", "Max group", "p99", "p99.9"],
        col_rows,
        [1.05, 1.2, 1.0, 1.05, 1.05, 0.9, 0.6, 0.75],
    )

    capacities = [512, 512, 512, 256, 256]
    prefix_rows = []
    for depth in range(1, 6):
        row = [str(depth)]
        for name in ["random", "semantic_mean"]:
            a = runs[name]["analysis"]
            stat = a["prefix_stats"][depth - 1]
            cap = math.prod(capacities[:depth])
            row.extend(
                [
                    fmt_int(stat["unique"]),
                    pct(stat["unique_ratio"], 2),
                    pct(stat["unique"] / cap, 4),
                    fmt_int(stat["max_group_size"]),
                ]
            )
        prefix_rows.append(row)
    add_table(
        doc,
        [
            "深度",
            "R used",
            "R unique/item",
            "R CUR",
            "R max",
            "S used",
            "S unique/item",
            "S CUR",
            "S max",
        ],
        prefix_rows,
        [0.45, 0.85, 0.85, 0.75, 0.65, 0.85, 0.85, 0.75, 0.65],
    )
    add_figure(doc, figures["prefix"], "图 4：prefix unique ratio。random 在 depth=3/4/5 的区分度略高。")
    add_figure(doc, figures["cur"], "图 5：分层 CUR。depth=1/2 对组合空间使用率最有解释力，后续理论空间远大于 item 数。")
    add_body(
        doc,
        "random 的 final unique ratio 为 98.61%，semantic_mean 为 97.24%。semantic_mean 的 collision 更高，直观原因是语义均值 embedding 会把同类、同标签、相似语义的 item 拉得更近；random 会注入更多 token-specific dispersion，因此更利于唯一性，但量化误差更高。",
    )

    add_heading(doc, "7. 近邻一致性与业务解释", 1)
    nn_rows = []
    for depth in range(1, 6):
        nn_rows.append(
            [
                str(depth),
                pct(runs["random"]["analysis"]["sampled_neighbor_prefix_agreement"]["agreement_by_depth"][str(depth)], 4),
                pct(runs["semantic_mean"]["analysis"]["sampled_neighbor_prefix_agreement"]["agreement_by_depth"][str(depth)], 4),
            ]
        )
    add_table(doc, ["Prefix depth", "random", "semantic_mean"], nn_rows, [1.3, 2.6, 2.6])
    add_body(
        doc,
        "近邻 prefix agreement 是在 3 万抽样 item 上，用 projected embedding 的 cosine 近邻计算的。depth=1 约 46%-47%，说明第一层 coarse cell 与连续空间近邻有一定一致性；depth>=3 很低，这是高基数 SID 的正常现象，也说明不能把 prefix 完全相同作为唯一相似性判断条件。",
    )

    add_heading(doc, "8. 同款不同颜色 Case Study", 1)
    add_body(
        doc,
        "针对 Manic Panic 染发剂样例，多个 item 在 cat_ids 和 flags 上完全一致，labels 也高度相似，但颜色差异会强烈影响 title/image embedding。比如 103189 与 103194 的 projected cosine 达到 0.960，但第一层 KMeans 仍可能分到不同 cell，这是硬聚类边界造成的离散跳变。当前 SID 更像 item-level 语义压缩特征，并不保证 prefix 表达 SPU/同款层级。"
    )
    add_table(
        doc,
        ["item_id", "semantic SID", "title cos vs 103189", "image cos", "projected cos"],
        [
            ["103189", "457-112-397-195-67", "1.000", "1.000", "1.000"],
            ["103192", "457-112-315-122-222", "0.926", "0.800", "0.842"],
            ["103191", "457-424-163-4-131", "0.951", "0.804", "0.819"],
            ["103196", "457-112-146-62-161", "0.959", "0.889", "0.940"],
            ["103194", "241-412-381-4-45", "0.966", "0.865", "0.960"],
            ["103187", "501-112-446-54-3", "0.887", "0.689", "0.683"],
            ["103181", "203-98-450-4-248", "0.909", "0.461", "0.658"],
        ],
        [0.9, 1.75, 1.25, 1.05, 1.15],
    )
    add_callout(
        doc,
        "业务含义",
        "如果目标是让同款不同颜色共享 coarse prefix，需要额外引入 parent_item_id/SPU/variation group、去颜色词 title embedding，或者使用 group mean embedding 训练前几层。纯无监督 RQ-OPQ 不会自动学习这个业务层级。",
    )

    add_heading(doc, "9. 结论与建议", 1)
    add_bullet(doc, "random：collision 更低、运行更快、SID 唯一性更强，适合作为保守的下游 categorical feature baseline。")
    add_bullet(doc, "semantic_mean：量化误差更低、重建 cosine 略好，语义聚合更强，但 collision 和 top group 更大。")
    add_bullet(doc, "当前两种结果都可以进入下游模型对比；最终选择不应只看 SID 指标，要看 CTR/排序模型的离线增益和线上稳定性。")
    add_bullet(doc, "如果只保留一个全量版本先试，我建议先用 random；如果模型能从语义聚合中受益，再引入 semantic_mean 或混合特征。")
    add_bullet(doc, "下一步可试 RQ L=4：dense concat 维度仍为 4*32 + 2*16 = 160，低于 200，且更直接降低 collision。")
    add_bullet(doc, "另一个方向是 semantic_mean 降低 cat/label/image 权重，缓解过度聚合带来的大 collision group。")

    add_heading(doc, "10. 产物索引", 1)
    add_table(
        doc,
        ["产物", "路径/说明"],
        [
            ["random metrics", "multimodal_rqopq_full_random/metrics.json"],
            ["random analysis", "multimodal_rqopq_full_random/analysis_summary.json"],
            ["random SID group", "multimodal_rqopq_full_random/sid_item_index/sid_groups_summary.csv"],
            ["semantic metrics", "multimodal_rqopq_full_semantic/metrics.json"],
            ["semantic analysis", "multimodal_rqopq_full_semantic/analysis_summary.json"],
            ["semantic SID group", "multimodal_rqopq_full_semantic/sid_item_index/sid_groups_summary.csv"],
            ["报告生成脚本", "create_full_rqopq_comparison_report_docx.py"],
        ],
        [1.7, 4.8],
    )

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
