#!/usr/bin/env python3
"""Create a comprehensive DOCX report for full random vs semantic_mean RQ-OPQ."""

from __future__ import annotations

import csv
import json
import math
from datetime import date
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sklearn.neighbors import NearestNeighbors


ROOT = Path(".")
RUNS = {
    "random": ROOT / "multimodal_rqopq_full_random",
    "semantic_mean": ROOT / "multimodal_rqopq_full_semantic",
}
OUT = ROOT / "full_rqopq_comprehensive_report.docx"
FIG_DIR = ROOT / "full_rqopq_comprehensive_figures"


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def dump_json(path: Path, obj: Any) -> None:
    path.write_text(json.dumps(obj, indent=2), encoding="utf-8")


def pct(x: float, digits: int = 2) -> str:
    return f"{x * 100:.{digits}f}%"


def fmt(x: float, digits: int = 4) -> str:
    return f"{x:.{digits}f}"


def fmt_int(x: int | float) -> str:
    return f"{int(x):,}"


def l2_normalize(x: np.ndarray) -> np.ndarray:
    return x / np.maximum(np.linalg.norm(x, axis=1, keepdims=True), 1e-12)


def entropy_from_counts(counts: np.ndarray) -> float:
    positive = counts[counts > 0].astype(np.float64)
    p = positive / positive.sum()
    return float(-(p * np.log2(p)).sum()) if positive.size else 0.0


def gini_from_counts(counts: np.ndarray) -> float:
    x = np.sort(counts.astype(np.float64))
    total = x.sum()
    if total <= 0:
        return 0.0
    n = x.size
    ranks = np.arange(1, n + 1, dtype=np.float64)
    return float((2.0 * np.sum(ranks * x)) / (n * total) - (n + 1.0) / n)


def rec_key(rec: dict[str, Any], key: str) -> float:
    aliases = {
        "per_dim_mse_mean": ["per_dim_mse_mean", "mse_mean"],
        "per_dim_mse_p95": ["per_dim_mse_p95", "mse_p95"],
        "per_dim_mse_p99": ["per_dim_mse_p99", "mse_p99"],
    }
    for k in aliases.get(key, [key]):
        if k in rec:
            return float(rec[k])
    raise KeyError(key)


def reconstruct(codes: np.ndarray, rq: np.ndarray, opq: np.ndarray, rotation: np.ndarray) -> np.ndarray:
    rq_levels = rq.shape[0]
    subspaces = opq.shape[0]
    subdim = opq.shape[2]
    out = np.zeros((codes.shape[0], rq.shape[2]), dtype=np.float32)
    for level in range(rq_levels):
        out += rq[level, codes[:, level]]
    rotated_residual = np.zeros_like(out)
    for subspace in range(subspaces):
        s0 = subspace * subdim
        s1 = s0 + subdim
        rotated_residual[:, s0:s1] = opq[subspace, codes[:, rq_levels + subspace]]
    out += rotated_residual @ rotation.T
    return out


def longest_common_prefix(a: np.ndarray, b: np.ndarray) -> np.ndarray:
    matches = a == b
    return np.cumprod(matches.astype(np.int8), axis=1).sum(axis=1)


def read_collision_group_sizes(path: Path) -> np.ndarray:
    sizes = []
    with path.open(newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            sizes.append(int(row["group_size"]))
    return np.asarray(sizes, dtype=np.int32)


def code_usage_rows(metrics: dict[str, Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for level in metrics["rq_kmeans"]["per_level"]:
        counts = np.asarray(level["counts"], dtype=np.int64)
        entropy = entropy_from_counts(counts)
        rows.append(
            {
                "name": f"RQ L{level['level'] + 1}",
                "k": counts.size,
                "cur": int((counts > 0).sum()) / counts.size,
                "entropy_ratio": entropy / math.log2(counts.size),
                "effective_ratio": (2.0**entropy) / counts.size,
                "gini": gini_from_counts(counts),
                "min": int(counts[counts > 0].min()),
                "p50": int(np.quantile(counts[counts > 0], 0.5)),
                "p95": int(np.quantile(counts[counts > 0], 0.95)),
                "p99": int(np.quantile(counts[counts > 0], 0.99)),
                "max": int(counts.max()),
                "counts": counts.tolist(),
            }
        )
    for sub in metrics["opq"]["history"][-1]["per_subspace"]:
        counts = np.asarray(sub["counts"], dtype=np.int64)
        entropy = entropy_from_counts(counts)
        rows.append(
            {
                "name": f"OPQ {sub['subspace']}",
                "k": counts.size,
                "cur": int((counts > 0).sum()) / counts.size,
                "entropy_ratio": entropy / math.log2(counts.size),
                "effective_ratio": (2.0**entropy) / counts.size,
                "gini": gini_from_counts(counts),
                "min": int(counts[counts > 0].min()),
                "p50": int(np.quantile(counts[counts > 0], 0.5)),
                "p95": int(np.quantile(counts[counts > 0], 0.95)),
                "p99": int(np.quantile(counts[counts > 0], 0.99)),
                "max": int(counts.max()),
                "counts": counts.tolist(),
            }
        )
    return rows


def compute_rich(run_dir: Path, sample_size: int = 60000, nn_sample_size: int = 12000) -> dict[str, Any]:
    out_path = run_dir / "rich_eval_summary.json"
    arrays_path = run_dir / "rich_eval_arrays.npz"
    if out_path.exists() and arrays_path.exists():
        return load_json(out_path)

    rng = np.random.default_rng(20260518)
    codes = np.load(run_dir / "sid_codes.npy", mmap_mode="r")
    projected = np.load(run_dir / "projected.npy", mmap_mode="r")
    rq = np.load(run_dir / "rq_codebooks.npy")
    opq = np.load(run_dir / "opq_codebooks.npy")
    rotation = np.load(run_dir / "opq_rotation.npy")
    n = codes.shape[0]

    sample_idx = np.sort(rng.choice(n, size=min(sample_size, n), replace=False))
    sample_codes = np.asarray(codes[sample_idx])
    sample_projected = np.asarray(projected[sample_idx], dtype=np.float32)
    sample_recon = reconstruct(sample_codes, rq, opq, rotation)
    per_dim_mse = ((sample_projected - sample_recon) ** 2).mean(axis=1)
    cos = np.einsum("ij,ij->i", sample_projected, sample_recon) / np.maximum(
        np.linalg.norm(sample_projected, axis=1) * np.linalg.norm(sample_recon, axis=1),
        1e-12,
    )

    nn_idx = np.sort(rng.choice(n, size=min(nn_sample_size, n), replace=False))
    nn_codes = np.asarray(codes[nn_idx])
    nn_projected = np.asarray(projected[nn_idx], dtype=np.float32)
    nn_recon = reconstruct(nn_codes, rq, opq, rotation)
    topk = 20
    original_norm = l2_normalize(nn_projected)
    recon_norm = l2_normalize(nn_recon)
    original_nn = NearestNeighbors(n_neighbors=topk + 1, metric="cosine", algorithm="brute")
    recon_nn = NearestNeighbors(n_neighbors=topk + 1, metric="cosine", algorithm="brute")
    original_nn.fit(original_norm)
    recon_nn.fit(recon_norm)
    original_indices = original_nn.kneighbors(original_norm, return_distance=False)[:, 1:]
    recon_indices = recon_nn.kneighbors(recon_norm, return_distance=False)[:, 1:]
    lcp = longest_common_prefix(
        np.repeat(nn_codes, topk, axis=0), nn_codes[original_indices.reshape(-1)]
    )
    random_neighbors = rng.integers(0, nn_codes.shape[0], size=nn_codes.shape[0] * topk)
    random_lcp = longest_common_prefix(
        np.repeat(nn_codes, topk, axis=0), nn_codes[random_neighbors]
    )
    recalls = np.empty(nn_codes.shape[0], dtype=np.float32)
    for i, (orig, rec) in enumerate(zip(original_indices, recon_indices, strict=True)):
        recalls[i] = len(set(orig.tolist()) & set(rec.tolist())) / topk

    # PCA 2-D scatter for sampled projected vectors.
    scatter_n = min(12000, sample_projected.shape[0])
    scatter_idx = np.sort(rng.choice(sample_projected.shape[0], size=scatter_n, replace=False))
    x = sample_projected[scatter_idx]
    centered = x - x.mean(axis=0, keepdims=True)
    _, _, vt = np.linalg.svd(centered, full_matrices=False)
    xy = centered @ vt[:2].T
    scatter_codes = sample_codes[scatter_idx, 0]

    np.savez_compressed(
        arrays_path,
        per_dim_mse=per_dim_mse.astype(np.float32),
        cosine=cos.astype(np.float32),
        neighbor_lcp=lcp.astype(np.int16),
        random_lcp=random_lcp.astype(np.int16),
        nn_recall=recalls.astype(np.float32),
        scatter_xy=xy.astype(np.float32),
        scatter_code=scatter_codes.astype(np.int32),
    )
    summary = {
        "sample_size": int(sample_idx.size),
        "nn_sample_size": int(nn_idx.size),
        "topk": topk,
        "per_dim_mse": {
            "mean": float(per_dim_mse.mean()),
            "p50": float(np.quantile(per_dim_mse, 0.5)),
            "p95": float(np.quantile(per_dim_mse, 0.95)),
            "p99": float(np.quantile(per_dim_mse, 0.99)),
        },
        "cosine": {
            "mean": float(cos.mean()),
            "p50": float(np.quantile(cos, 0.5)),
            "p05": float(np.quantile(cos, 0.05)),
            "p01": float(np.quantile(cos, 0.01)),
        },
        "neighbor_lcp": {
            "mean": float(lcp.mean()),
            "random_mean": float(random_lcp.mean()),
            "hist": np.bincount(lcp, minlength=codes.shape[1] + 1).astype(int).tolist(),
            "random_hist": np.bincount(random_lcp, minlength=codes.shape[1] + 1).astype(int).tolist(),
        },
        "nn_recall_at_20": {
            "mean": float(recalls.mean()),
            "p50": float(np.quantile(recalls, 0.5)),
            "p90": float(np.quantile(recalls, 0.9)),
            "p95": float(np.quantile(recalls, 0.95)),
        },
    }
    dump_json(out_path, summary)
    return summary


def load_all() -> dict[str, dict[str, Any]]:
    runs = {}
    for name, run_dir in RUNS.items():
        runs[name] = {
            "dir": run_dir,
            "metrics": load_json(run_dir / "metrics.json"),
            "analysis": load_json(run_dir / "analysis_summary.json"),
            "rich": compute_rich(run_dir),
            "usage": None,
        }
        runs[name]["usage"] = code_usage_rows(runs[name]["metrics"])
    return runs


def save_common_figures(runs: dict[str, dict[str, Any]]) -> dict[str, Path]:
    FIG_DIR.mkdir(exist_ok=True)
    plt.style.use("seaborn-v0_8-whitegrid")
    colors = {"random": "#4C78A8", "semantic_mean": "#F58518"}
    figs: dict[str, Path] = {}

    def save(name: str) -> Path:
        path = FIG_DIR / f"{name}.png"
        figs[name] = path
        plt.tight_layout()
        plt.savefig(path, dpi=180)
        plt.close()
        return path

    # residual curve
    plt.figure(figsize=(7.2, 4.1))
    stages = ["RQ1", "RQ2", "RQ3", "OPQ"]
    for name, run in runs.items():
        m = run["metrics"]
        vals = [x["residual_mse_after_level"] for x in m["rq_kmeans"]["per_level"]]
        vals.append(m["opq"]["final_original_space_mse"])
        plt.plot(stages, vals, marker="o", linewidth=2.2, label=name, color=colors[name])
    plt.title("Residual SSE Curve")
    plt.ylabel("SSE per item")
    plt.legend()
    save("01_residual_sse_curve")

    # codebook usage quality
    labels = ["RQ1", "RQ2", "RQ3", "OPQ0", "OPQ1"]
    x = np.arange(len(labels))
    width = 0.36
    plt.figure(figsize=(7.2, 4.1))
    for offset, name in [(-width / 2, "random"), (width / 2, "semantic_mean")]:
        vals = [row["entropy_ratio"] * 100 for row in runs[name]["usage"]]
        plt.bar(x + offset, vals, width, label=name, color=colors[name])
    plt.ylim(94, 100.4)
    plt.xticks(x, labels)
    plt.ylabel("entropy / max entropy (%)")
    plt.title("Codebook Usage Balance")
    plt.legend()
    save("02_codebook_usage_balance")

    # sorted code frequency
    for name, run in runs.items():
        plt.figure(figsize=(7.2, 4.1))
        for row in run["usage"]:
            counts = np.sort(np.asarray(row["counts"], dtype=np.float64))[::-1]
            plt.plot(np.arange(1, counts.size + 1), counts, linewidth=1.8, label=row["name"])
        plt.title(f"Sorted Code Usage Frequency - {name}")
        plt.xlabel("code rank")
        plt.ylabel("items")
        plt.legend(fontsize=8)
        save(f"03_sorted_code_usage_{name}")

    # reconstruction histograms
    for arr_name, title, key in [
        ("cosine", "Reconstruction Cosine Histogram", "cosine"),
        ("per_dim_mse", "Per-dim Reconstruction MSE Histogram", "per_dim_mse"),
        ("nn_recall", "Reconstruction NN Recall@20 Histogram", "nn_recall"),
    ]:
        plt.figure(figsize=(7.2, 4.1))
        for name, run in runs.items():
            arrays = np.load(run["dir"] / "rich_eval_arrays.npz")
            data = arrays[key]
            plt.hist(data, bins=50, alpha=0.55, density=True, label=name, color=colors[name])
        plt.title(title)
        plt.legend()
        save(f"04_{arr_name}_hist")

    # collision group size histogram
    plt.figure(figsize=(7.2, 4.1))
    bins = np.array([2, 3, 4, 5, 6, 8, 10, 20, 50, 100, 300, 1000, 3000])
    for name, run in runs.items():
        sizes = read_collision_group_sizes(run["dir"] / "sid_item_index" / "sid_groups_summary.csv")
        plt.hist(sizes, bins=bins, alpha=0.55, label=name, color=colors[name])
    plt.xscale("log")
    plt.yscale("log")
    plt.title("Collision Group Size Histogram")
    plt.xlabel("group size")
    plt.ylabel("groups")
    plt.legend()
    save("05_collision_group_size_hist")

    # prefix uniqueness and hierarchical CUR
    depths = np.arange(1, 6)
    plt.figure(figsize=(7.2, 4.1))
    for name, run in runs.items():
        vals = [x["unique_ratio"] * 100 for x in run["analysis"]["prefix_stats"]]
        plt.plot(depths, vals, marker="o", linewidth=2.2, label=name, color=colors[name])
    plt.xticks(depths)
    plt.title("Prefix Unique Ratio")
    plt.xlabel("SID depth")
    plt.ylabel("unique prefix / item (%)")
    plt.legend()
    save("06_prefix_unique_ratio")

    capacities = [512, 512, 512, 256, 256]
    plt.figure(figsize=(7.2, 4.1))
    for name, run in runs.items():
        cap = 1
        vals = []
        for stat, c in zip(run["analysis"]["prefix_stats"], capacities, strict=True):
            cap *= c
            vals.append(stat["unique"] / cap * 100)
        plt.plot(depths, vals, marker="o", linewidth=2.2, label=name, color=colors[name])
    plt.yscale("log")
    plt.xticks(depths)
    plt.title("Hierarchical CUR")
    plt.xlabel("SID depth")
    plt.ylabel("used prefix / theoretical space (%)")
    plt.legend()
    save("07_hierarchical_cur")

    # neighbor LCP histogram
    for name, run in runs.items():
        arrays = np.load(run["dir"] / "rich_eval_arrays.npz")
        hist = np.bincount(arrays["neighbor_lcp"], minlength=6)
        random_hist = np.bincount(arrays["random_lcp"], minlength=6)
        plt.figure(figsize=(7.2, 4.1))
        w = 0.35
        xs = np.arange(6)
        plt.bar(xs - w / 2, hist / hist.sum() * 100, width=w, label="nearest neighbors", color=colors[name])
        plt.bar(xs + w / 2, random_hist / random_hist.sum() * 100, width=w, label="random pairs", color="#999999")
        plt.title(f"Neighbor Prefix LCP Histogram - {name}")
        plt.xlabel("longest common prefix length")
        plt.ylabel("pair share (%)")
        plt.xticks(xs)
        plt.legend()
        save(f"08_neighbor_lcp_{name}")

    # scatter
    for name, run in runs.items():
        arrays = np.load(run["dir"] / "rich_eval_arrays.npz")
        xy = arrays["scatter_xy"]
        code = arrays["scatter_code"]
        plt.figure(figsize=(6.4, 5.1))
        plt.scatter(xy[:, 0], xy[:, 1], c=code, cmap="tab20", s=3, alpha=0.6, linewidths=0)
        plt.title(f"PCA Scatter Colored by First SID Token - {name}")
        plt.xticks([])
        plt.yticks([])
        save(f"09_pca_scatter_{name}")

    return figs


def set_run_font(run, size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, size=10.6, color="222222", after=5, line_spacing=1.08) -> None:
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=100, bottom=70, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B7C3D1") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        old = borders.find(qn(f"w:{edge}"))
        if old is not None:
            borders.remove(old)
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)


def setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.6)
    for name, size, color in [
        ("Title", 21, "0B2545"),
        ("Heading 1", 15, "2E74B5"),
        ("Heading 2", 12.5, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    return doc


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_run_font(run, color="2E74B5" if level <= 2 else "1F4D78")


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    style_paragraph(p)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    style_paragraph(p, after=3, line_spacing=1.12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for c, h in zip(table.rows[0].cells, headers, strict=True):
        c.text = h
    for row in rows:
        cells = table.add_row().cells
        for c, v in zip(cells, row, strict=True):
            c.text = v
    table.autofit = False
    for i, row in enumerate(table.rows):
        for c, w in zip(row.cells, widths, strict=True):
            c.width = Inches(w)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c)
            if i == 0:
                shade_cell(c, "F2F4F7")
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.03
                for r in p.runs:
                    set_run_font(r, size=8.3, bold=(i == 0), color="1F3349" if i == 0 else "222222")
    set_table_borders(table)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_margins(cell, top=130, bottom=130, start=170, end=170)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=10.8, bold=True, color="1F3A5F")
    p.add_run("  ")
    r = p.add_run(text)
    set_run_font(r, size=10.2, color="222222")
    style_paragraph(p, after=0)
    set_table_borders(table, color="D8DEE8")
    doc.add_paragraph()


def add_fig(doc: Document, path: Path, caption: str, width=6.65) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(7)
    for r in cap.runs:
        set_run_font(r, size=8.8, italic=True, color="536173")


def main() -> None:
    runs = load_all()
    figs = save_common_figures(runs)
    doc = setup_doc()
    title = doc.add_paragraph(style="Title")
    title.add_run("全量商品 RQ-OPQ SID 综合评估报告：random vs semantic_mean")
    style_paragraph(title, size=21, color="0B2545", after=3, line_spacing=1.05)
    sub = doc.add_paragraph(f"生成日期：{date.today().isoformat()} | 样本：11,168,159 items | SID: RQ K=512,L=3 + OPQ M=2,K=256")
    style_paragraph(sub, size=9.5, color="536173", after=10)

    add_callout(
        doc,
        "总体结论",
        "random 更稳、更快、collision 更低；semantic_mean 的量化重建质量明显更好，但会把语义相近 item 聚得更紧，导致 collision 与 top group 更大。两者都没有 codebook collapse，建议都进入下游模型实验，最终以排序/CTR 特征增益决策。",
    )

    add_heading(doc, "1. 核心指标总览")
    rows = []
    for name, run in runs.items():
        m, a, rich = run["metrics"], run["analysis"], run["rich"]
        rows.append([
            name,
            fmt_int(m["collisions"]["unique_sid"]),
            pct(m["collisions"]["collision_rate_extra"], 3),
            pct(a["collision_distribution"]["collision_items"] / a["num_items"], 3),
            fmt_int(m["collisions"]["max_collision_group_size"]),
            fmt(m["opq"]["final_original_space_mse"], 5),
            fmt(rich["cosine"]["mean"], 4),
            f"{m['elapsed_seconds']/60:.1f}m",
        ])
    add_table(doc, ["模式", "Unique SID", "Extra collision", "Collision items", "Max group", "OPQ SSE", "Cos mean", "耗时"], rows, [0.95, 1.05, 1.0, 1.0, 0.8, 0.85, 0.8, 0.65])
    add_fig(doc, figs["01_residual_sse_curve"], "图 1：Residual SSE curve。semantic_mean 在 RQ 与 OPQ 每个阶段误差都更低。")

    add_heading(doc, "2. 指标口径")
    add_table(
        doc,
        ["指标", "计算方式", "作用"],
        [
            ["Codebook CUR", "used code / K", "检查单层码本是否塌缩"],
            ["Entropy / Effective ratio", "entropy(counts)/log2(K), 2^entropy/K", "检查码字使用是否均匀"],
            ["Sorted code usage", "每列 code count 降序曲线", "观察头部 code 是否过度集中"],
            ["Hierarchical CUR", "unique prefix / theoretical prefix space", "观察分层组合空间使用率"],
            ["Unique prefix ratio", "unique prefix / item count", "观察前缀区分 item 的能力"],
            ["Extra collision rate", "(N - unique SID) / N", "最终 SID 额外冲突率"],
            ["Collision item coverage", "非 singleton group 中 item / N", "有多少 item 处于 collision group"],
            ["Reconstruction cosine", "cos(projected, reconstructed)", "SID 重建向量和原 PCA 向量一致性"],
            ["NN Recall@20", "原向量 top20 近邻与重建向量 top20 近邻重合率", "重建后近邻结构保留程度"],
            ["Neighbor LCP", "连续空间近邻的 SID 最长公共前缀分布", "SID 前缀是否承载近邻结构"],
        ],
        [1.3, 2.35, 2.65],
    )

    add_heading(doc, "3. Codebook 使用质量")
    usage_rows = []
    for name, run in runs.items():
        for row in run["usage"]:
            usage_rows.append([
                name,
                row["name"],
                pct(row["cur"]),
                pct(row["entropy_ratio"]),
                pct(row["effective_ratio"]),
                fmt(row["gini"], 3),
                f"{fmt_int(row['min'])}/{fmt_int(row['p50'])}/{fmt_int(row['p99'])}/{fmt_int(row['max'])}",
            ])
    add_table(doc, ["模式", "列", "CUR", "Entropy", "Effective", "Gini", "Min/P50/P99/Max"], usage_rows, [0.85, 0.75, 0.75, 0.85, 0.85, 0.55, 1.9])
    add_fig(doc, figs["02_codebook_usage_balance"], "图 2：Codebook usage balance。两种模式的 entropy ratio 都接近 1，无码本塌缩。")
    add_fig(doc, figs["03_sorted_code_usage_random"], "图 3：random sorted code usage frequency。曲线越平，码字越均匀。")
    add_fig(doc, figs["03_sorted_code_usage_semantic_mean"], "图 4：semantic_mean sorted code usage frequency。RQ1 会比后续层更不均衡，但整体仍健康。")

    add_heading(doc, "4. 重建质量与近邻结构")
    rec_rows = []
    for name, run in runs.items():
        rich = run["rich"]
        rec_rows.append([
            name,
            fmt(rich["per_dim_mse"]["mean"], 6),
            fmt(rich["per_dim_mse"]["p95"], 6),
            fmt(rich["cosine"]["mean"], 4),
            fmt(rich["cosine"]["p05"], 4),
            fmt(rich["nn_recall_at_20"]["mean"], 4),
            fmt(rich["nn_recall_at_20"]["p90"], 4),
        ])
    add_table(doc, ["模式", "MSE mean", "MSE p95", "Cos mean", "Cos p05", "NN R@20", "NN R@20 p90"], rec_rows, [0.95, 0.95, 0.85, 0.85, 0.8, 0.85, 0.85])
    add_fig(doc, figs["04_cosine_hist"], "图 5：Reconstruction cosine histogram。semantic_mean 分布略向高 cosine 侧移动。")
    add_fig(doc, figs["04_per_dim_mse_hist"], "图 6：Per-dim reconstruction MSE histogram。semantic_mean 的误差分布更低。")
    add_fig(doc, figs["04_nn_recall_hist"], "图 7：Reconstruction NN Recall@20 histogram。衡量重建后近邻集合保留程度。")

    add_heading(doc, "5. Collision、Prefix 与分层 CUR")
    col_rows = []
    for name, run in runs.items():
        m, a = run["metrics"], run["analysis"]
        dist = a["collision_distribution"]
        col_rows.append([
            name,
            fmt_int(m["collisions"]["collision_groups"]),
            fmt_int(dist["collision_items"]),
            fmt(dist["mean_collision_group_size"], 2),
            fmt(dist["quantiles"]["0.95"], 0),
            fmt(dist["quantiles"]["0.99"], 0),
            fmt(dist["quantiles"]["0.999"], 1),
            fmt_int(dist["max_group_size"]),
        ])
    add_table(doc, ["模式", "Groups", "Items", "Mean size", "P95", "P99", "P99.9", "Max"], col_rows, [0.95, 0.9, 0.95, 0.8, 0.6, 0.6, 0.65, 0.75])
    add_fig(doc, figs["05_collision_group_size_hist"], "图 8：Collision group size histogram。semantic_mean 的大 group 更多，random 更保守。")
    add_fig(doc, figs["06_prefix_unique_ratio"], "图 9：Prefix unique ratio。depth=3 后开始具备主要区分力，最终 random 更高。")
    add_fig(doc, figs["07_hierarchical_cur"], "图 10：Hierarchical CUR。depth=1/2 更适合看组合空间利用，后续理论空间过大。")

    prefix_rows = []
    caps = [512, 512, 512, 256, 256]
    for depth in range(1, 6):
        row = [str(depth)]
        for name in ["random", "semantic_mean"]:
            stat = runs[name]["analysis"]["prefix_stats"][depth - 1]
            cap = math.prod(caps[:depth])
            row += [fmt_int(stat["unique"]), pct(stat["unique_ratio"], 2), pct(stat["unique"] / cap, 4), fmt_int(stat["max_group_size"])]
        prefix_rows.append(row)
    add_table(doc, ["深度", "R used", "R uniq/item", "R CUR", "R max", "S used", "S uniq/item", "S CUR", "S max"], prefix_rows, [0.42, 0.82, 0.82, 0.72, 0.6, 0.82, 0.82, 0.72, 0.6])

    add_heading(doc, "6. Neighbor Prefix / LCP 诊断")
    lcp_rows = []
    for name, run in runs.items():
        rich = run["rich"]["neighbor_lcp"]
        hist = np.asarray(rich["hist"], dtype=np.float64)
        hist = hist / hist.sum()
        lcp_rows.append([name, fmt(rich["mean"], 4), fmt(rich["random_mean"], 4)] + [pct(v, 3) for v in hist.tolist()])
    add_table(doc, ["模式", "Neighbor mean LCP", "Random mean LCP", "LCP0", "LCP1", "LCP2", "LCP3", "LCP4", "LCP5"], lcp_rows, [0.8, 1.05, 1.0, 0.65, 0.65, 0.65, 0.65, 0.65, 0.65])
    add_fig(doc, figs["08_neighbor_lcp_random"], "图 11：random 的连续空间近邻 LCP 分布，与随机 pair 对比。")
    add_fig(doc, figs["08_neighbor_lcp_semantic_mean"], "图 12：semantic_mean 的连续空间近邻 LCP 分布，与随机 pair 对比。")
    add_body(doc, "LCP 指标的意义是：如果连续空间近邻也更容易共享 SID 前缀，说明 SID 的层级结构承载了部分语义邻近关系。当前 depth=1 的一致性明显高于随机，但 depth 较深后完全相同的比例很低，这是高基数 item-level SID 的正常现象。")

    add_heading(doc, "7. 结构可视化")
    add_fig(doc, figs["09_pca_scatter_random"], "图 13：random projected PCA scatter，按第一层 SID token 着色。")
    add_fig(doc, figs["09_pca_scatter_semantic_mean"], "图 14：semantic_mean projected PCA scatter，按第一层 SID token 着色。")
    add_body(doc, "散点图是 12k 抽样 projected embedding 的二维 PCA 投影。它不能替代下游效果评估，但可以帮助观察第一层 SID 是否在连续空间中形成局部结构。")

    add_heading(doc, "8. 业务 Case 与解释")
    add_body(doc, "Manic Panic 染发剂样例显示，同款不同颜色的 item 即使 cat_ids 完全一致，也可能因为 title/image 中颜色信号很强而落入不同第一层 SID。该现象不是 itemid 对齐错误，而是无监督 RQ-OPQ 的 item-level 压缩目标与 SPU/同款业务层级不完全一致。")
    add_table(
        doc,
        ["item_id", "semantic SID", "title cos", "image cos", "projected cos"],
        [
            ["103189", "457-112-397-195-67", "1.000", "1.000", "1.000"],
            ["103192", "457-112-315-122-222", "0.926", "0.800", "0.842"],
            ["103191", "457-424-163-4-131", "0.951", "0.804", "0.819"],
            ["103196", "457-112-146-62-161", "0.959", "0.889", "0.940"],
            ["103194", "241-412-381-4-45", "0.966", "0.865", "0.960"],
            ["103187", "501-112-446-54-3", "0.887", "0.689", "0.683"],
            ["103181", "203-98-450-4-248", "0.909", "0.461", "0.658"],
        ],
        [0.85, 1.7, 0.85, 0.85, 0.9],
    )

    add_heading(doc, "9. 建议")
    add_bullet(doc, "当前先保留两套结果进入下游模型实验；random 作为唯一性更稳的 baseline，semantic_mean 作为重建/语义质量更强的版本。")
    add_bullet(doc, "如果只先选一套给业务试用，random 更保守；如果后续模型更依赖语义聚合特征，可以加入 semantic_mean 或做双 SID 特征。")
    add_bullet(doc, "若业务希望同款不同颜色共享 coarse prefix，需要引入 parent_item_id/SPU/variation group、去颜色 title，或者用 group mean embedding 训练前几层。")
    add_bullet(doc, "下一组实验建议优先尝试 RQ L=4，dense concat 维度仍为 160，低于 200，且更直接降低 collision。")

    add_heading(doc, "10. 产物")
    add_table(
        doc,
        ["产物", "路径"],
        [
            ["综合报告", OUT.as_posix()],
            ["综合图表目录", FIG_DIR.as_posix()],
            ["random rich eval", "multimodal_rqopq_full_random/rich_eval_summary.json"],
            ["semantic rich eval", "multimodal_rqopq_full_semantic/rich_eval_summary.json"],
            ["random SID group", "multimodal_rqopq_full_random/sid_item_index/sid_groups_summary.csv"],
            ["semantic SID group", "multimodal_rqopq_full_semantic/sid_item_index/sid_groups_summary.csv"],
        ],
        [1.6, 4.9],
    )

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
